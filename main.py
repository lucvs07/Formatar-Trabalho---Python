# Formatação de um documento em abnt
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION




# Função para Configurar as margens do documento
def config_margins(section, top, bottom, left, right):
    section.top_margin = Pt(top)
    section.bottom_margin = Pt(bottom)
    section.left_margin = Pt(left)
    section.right_margin = Pt(right)
    
# Função configurar parágrafo
def config_paragraph(object_paragraph, size):
    font = object_paragraph.font
    font.name = 'Arial'
    font.size = Pt(size)
    
    paragraph_format = object_paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    

# Criar parágrafo
def run_paragraph(text, object_text):
    object_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = object_text.add_run(text)

# Adicionar título 
def create_title(object_title, text):
    object_title_run = object_title.add_run(text)

    

    
# Função para criar o documento
def create_document(text_title_trabalho, text_title_materia, text_data_equipe, alunos, professor):
    doc = Document('modelo.docx')
    
    # Definir margens
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    config_margins(section, 85.05, 56.7, 85.05, 56.7)
    
    # Página Equipe
    title_equipe_section = doc.add_heading(level=1)
    title_equipe_section_text = '1.Equipe'
    create_title(object_title=title_equipe_section, text=title_equipe_section_text)
    
    title_trabalho = doc.add_paragraph()
    title_trabalho_run = title_trabalho.add_run(text_title_trabalho)
    title_trabalho.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_trabalho_run.font.size = Pt(16)
    
    title_materia = doc.add_paragraph()
    title_materia_run = title_materia.add_run(text_title_materia)
    title_materia.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_materia_run.font.size = Pt(14)
    title_materia_run.font.bold = True
    
    data_equipe = doc.add_paragraph()
    data_equipe_run = data_equipe.add_run(text_data_equipe)
    data_equipe.alignment = WD_ALIGN_PARAGRAPH.LEFT
    data_equipe_run.font.size = Pt(12)
    data_equipe_run.font.bold = True
    
    space = doc.add_paragraph()
        
    title_alunos = doc.add_heading(level=2)
    title_alunos_text = 'Alunos Responsáveis'
    create_title(object_title=title_alunos, text=title_alunos_text)
    for aluno in alunos :
        nome = doc.add_paragraph(style='List Bullet')
        nome_text = aluno
        run_paragraph(text=nome_text, object_text=nome)
        aluno_text = doc.add_paragraph()
        aluno_text_run = aluno_text.add_run('Estudante de Engenharia da Computação')
        aluno_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
        aluno_text_run.font.size = Pt(12)
        aluno_text_run.font.italic = True
    
    space = doc.add_paragraph()
    
    title_professor = doc.add_heading(level=2)
    title_professor_text = 'Professor Responsável'
    create_title(object_title=title_professor, text=title_professor_text)
    professor_sub = doc.add_paragraph(style='List Bullet')
    professor_sub_text = professor
    run_paragraph(text=professor_sub_text, object_text=professor_sub)
    materia_professor = doc.add_paragraph()
    materia_professor_run = materia_professor.add_run(f'Professor de {text_title_materia}')
    materia_professor.alignment = WD_ALIGN_PARAGRAPH.LEFT
    materia_professor_run.font.size = Pt(12)
    materia_professor_run.font.italic = True
    
    doc.add_page_break()
    
    
    
    
    # Adicionando Títulos
    title2 = doc.add_heading(level=2)
    title2_text = 'Drone Multirotor'
    create_title(object_title=title2, text=title2_text)

    
    # Definir estilo de parágrafo
    style = doc.styles['Normal']
    config_paragraph(object_paragraph=style, size=12)


    # Adicionar seção de texto (parágrafo normal)
        
    paragraph_1 = doc.add_paragraph()
    paragraph_1_text = 'O drone multirotor robusto e versátil como o DJI Matrice 600 Pro, foi escolhido, devido a suas características se adequarem para a detecção de manchas de óleo no oceano, como por exemplo : sua capacidade de carregar equipamentos pesados e avançados, como câmeras multiespectrais e sistemas de transmissão de dados.'
    run_paragraph(paragraph_1_text, object_text=paragraph_1)
    
    paragraph_2 = doc.add_paragraph()
    paragraph_2_text = 'Drones do tipo multirotor apresentam um ótimo controle que demonstra precisão e estabilidade, característica que se apresenta como fundamental para o desempenho  atividades de mapeamento e monitoramento.'
    run_paragraph(paragraph_2_text, object_text=paragraph_2)
    
    paragraph_3 = doc.add_paragraph()
    paragraph_3_text = 'Além disso, a capacidade de planar no ar e executar movimentos com exatidão é imprescindível para colher informações profundas e precisas acerca das manchas de óleo presentes no oceano.'
    run_paragraph(paragraph_3_text, object_text=paragraph_3)
    
    # Seção 2
    doc.add_page_break()
    title3 = doc.add_heading(level=2)
    create_title(object_title=title3, text='Componentes do Drone')
    
    
    
    
    
    # Salvar o documento
    doc.save(f'{text_title_trabalho}.docx')
titulo_trabalho = input('Insira o título do trabalho: ')
materia = input('Insira a matéria do trabalho: ')
data = input('Insira a data do trabalho no formato (ex: 29 de Março de 2024): ')
alunos = []
n = int(input('Quantos alunos fazem parte do seu grupo (contando com você): '))
i = 0
while i < n:
    nome = input(f'Insira o nome e o email no formato (nome | email) do aluno {i+1}: ')
    alunos.append(nome)
    i += 1
professor = input('Insira o nome do professor no formato (nome | email): ')
create_document(text_title_trabalho=titulo_trabalho, text_title_materia=materia, text_data_equipe=data, alunos=alunos, professor=professor)
print('Formatação Concluída')